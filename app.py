from flask import Flask, request, send_file, render_template
import pdfplumber
from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import os
import time

# Create the Flask application instance
app = Flask(__name__, static_folder='static')


def extract_details(pdf_path, word_template_path, output_word_path):
    doc = Document(word_template_path)

    # Function to set cell borders
    def set_cell_border(cell, **kwargs):
        tc = cell._element
        tcPr = tc.get_or_add_tcPr()

        for border_name in ['top', 'bottom', 'start', 'end', 'left', 'right', 'insideH', 'insideV']:
            if border_name in kwargs:
                border = OxmlElement(f"w:{border_name}")
                for key in kwargs[border_name]:
                    border.set(qn(f"w:{key}"), str(kwargs[border_name][key]))
                tcPr.append(border)

    # Extract product details
    def extract_product_details():
        page_width = 387
        top_bound = 240
        bottom_bound = 630

        product_col = (73, top_bound, page_width * 0.78, bottom_bound)
        model_col = (page_width * 0.78, top_bound, page_width * 0.90, bottom_bound)
        quantity_col = (page_width * 0.90, top_bound, page_width * 0.98, bottom_bound)

        stop_keyword = "CMISG Services"
        line_spacing_threshold = 10

        all_data = []
        keyword_found = False

        with pdfplumber.open(pdf_path) as pdf:
            for page in pdf.pages:
                if keyword_found:
                    break

                product_lines, model_lines, quantity_lines = [], [], []

                for word in page.extract_words():
                    x0, y0, x1, y1, text = word["x0"], word["top"], word["x1"], word["bottom"], word["text"]
                    if product_col[0] <= x0 <= product_col[2] and product_col[1] <= y0 <= product_col[3]:
                        product_lines.append((y0, text))
                    elif model_col[0] <= x0 <= model_col[2] and model_col[1] <= y0 <= model_col[3]:
                        model_lines.append((y0, text))
                    elif quantity_col[0] <= x0 <= quantity_col[2] and quantity_col[1] <= y0 <= quantity_col[3]:
                        quantity_lines.append((y0, text))

                product_lines.sort()
                model_lines.sort()
                quantity_lines.sort()

                def combine_into_cells(lines):
                    combined_cells = []
                    current_cell = ""
                    last_y = None

                    for y, text in lines:
                        if last_y is not None and abs(y - last_y) > line_spacing_threshold * 2:
                            combined_cells.append(current_cell.strip())
                            current_cell = text
                        else:
                            current_cell += " " + text
                        last_y = y

                    if current_cell:
                        combined_cells.append(current_cell.strip())
                    return combined_cells

                product_cells = combine_into_cells(product_lines)
                model_cells = combine_into_cells(model_lines)
                quantity_cells = combine_into_cells(quantity_lines)

                max_cells = max(len(product_cells), len(model_cells), len(quantity_cells))

                for i in range(max_cells):
                    if (i < len(product_cells) and stop_keyword.lower() in product_cells[i].lower()) or \
                       (i < len(model_cells) and stop_keyword.lower() in model_cells[i].lower()) or \
                       (i < len(quantity_cells) and stop_keyword.lower() in quantity_cells[i].lower()):
                        keyword_found = True
                        break

                    product_value = product_cells[i] if i < len(product_cells) else ''
                    model_value = model_cells[i] if i < len(model_cells) else ''
                    quantity_value = quantity_cells[i] if i < len(quantity_cells) else ''

                    all_data.append([product_value, model_value, quantity_value])

        return all_data

    # Extract customer details
    def extract_customer_details():
        page_width = 600  
        top_bound = 110
        bottom_bound = 210

        column_1 = (0, top_bound, page_width * 0.10, bottom_bound)
        column_2 = (page_width * 0.10, top_bound, page_width * 0.61, bottom_bound)
        column_3 = (page_width * 0.61, top_bound, page_width * 0.73, bottom_bound)
        column_4 = (page_width * 0.73, top_bound, page_width, bottom_bound)

        all_text_data = []

        with pdfplumber.open(pdf_path) as pdf:
            page = pdf.pages[0]

            col_1_text = page.within_bbox(column_1).extract_text() or ''
            col_2_text = page.within_bbox(column_2).extract_text() or ''
            col_3_text = page.within_bbox(column_3).extract_text() or ''
            col_4_text = page.within_bbox(column_4).extract_text() or ''

            col_1_lines = col_1_text.strip().split('\n')
            col_2_lines = col_2_text.strip().split('\n')
            col_3_lines = col_3_text.strip().split('\n')
            col_4_lines = col_4_text.strip().split('\n')

            max_len = max(len(col_1_lines), len(col_2_lines), len(col_3_lines), len(col_4_lines))
            col_1_lines += [''] * (max_len - len(col_1_lines))
            col_2_lines += [''] * (max_len - len(col_2_lines))
            col_3_lines += [''] * (max_len - len(col_3_lines))
            col_4_lines += [''] * (max_len - len(col_4_lines))

            for c1, c2, c3, c4 in zip(col_1_lines, col_2_lines, col_3_lines, col_4_lines):
                all_text_data.append([c1, c2, c3, c4])

        return all_text_data

    # Extract data from PDF
    product_data = extract_product_details()
    customer_data = extract_customer_details()

    # Add customer details to the document
    customer_details_mapping = {
        'CustomerContactPerson': customer_data[3][1],  
        'CustomerCompanyName': customer_data[0][1],    
        'Address': customer_data[1][1]                  
    }

    # Replace placeholders in the document for customer details
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for placeholder, value in customer_details_mapping.items():
                    placeholder_text = f"{{{{{placeholder}}}}}"
                    if placeholder_text in cell.text:
                        cell.text = cell.text.replace(placeholder_text, str(value))
                        
    # Add product details to a new table in the Word document
    if product_data:
        table = doc.tables[3]
        table.style = 'Table Grid'
        
        for row_data in product_data:
            if len(row_data) >= 3:
                row_cells = table.add_row().cells
                row_cells[2].text = row_data[1]  # Product value
                row_cells[3].text = row_data[0]  # Model value
                row_cells[4].text = row_data[2]  # Quantity value

                for idx in [2, 4]:  # Center both Product and Quantity
                    paragraph = row_cells[idx].paragraphs[0]
                    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

                # Set cell borders
                for cell in row_cells:
                    set_cell_border(cell,
                                    top={"sz": 5, "val": "single", "color": "000000"},
                                    bottom={"sz": 5, "val": "single", "color": "000000"},
                                    left={"sz": 5, "val": "single", "color": "000000"},
                                    right={"sz": 5, "val": "single", "color": "000000"})

    # Save the Word document
    doc.save(output_word_path)
    print("Extraction complete.")

@app.route('/')
def home():
    return render_template('upload.html')

@app.route('/upload', methods=['POST'])
def upload_file():
    # Check if both 'pdf_file' and 'word_template' files are in the request
    if 'pdf_file' not in request.files:
        return "PDF  file required", 400
    elif 'word_file' not in request.files:
        return "word file required", 400


    # file = request.files['file']
    pdf_file = request.files['pdf_file']
    word_file = request.files['word_file']
    
    if pdf_file.filename == '' or word_file == '':
        return "Please select both PDF and Word files", 400

    # Create an uploads directory if it doesn't exist
    if not os.path.exists('uploads'):
        os.makedirs('uploads')

    # Save the uploaded PDF temporarily
    pdf_path = os.path.join('uploads', pdf_file.filename)
    word_template_path = os.path.join('uploads', word_file.filename)
    pdf_file.save(pdf_path)
    word_file.save(word_template_path)

    # Define your Word template path
    # word_template_path = "CAF test.docx"

    # Create a unique output filename based on the current timestamp
    timestamp = int(time.time())  # Get the current time in seconds
    output_word_path = f"output_combined_{timestamp}.docx"

    # Extract details
    extract_details(pdf_path, word_template_path, output_word_path)

    # Clean up: remove the uploaded PDF after processing
    os.remove(pdf_path)
    os.remove(word_template_path)

    # Return the generated Word document
    return send_file(output_word_path, as_attachment=True)

if __name__ == '__main__':
    app.run(debug=True)
